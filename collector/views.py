from django.http import Http404, HttpResponse # Import HttpResponse
from wsgiref.util import FileWrapper # Efficiently stream large files (optional but good)
import mimetypes # To guess content type
import os # To work with file paths
import io
import csv
import zipfile
import datetime
import shutil
import tempfile
import time
from django.utils.timezone import now
from django.urls import reverse
from django.conf import settings
from collections import defaultdict
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import render, get_object_or_404
from django.http import FileResponse, Http404, HttpResponseServerError
from django.views import generic # Using generic class-based views for simplicity
from django.views.decorators.http import require_POST # For the export view
import json # To parse priorities from POST
from django.db.models.fields import files as file_fields
from rest_framework import generics, permissions
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .serializers import OplogEntrySerializer, TargetSerializer, CredentialSerializer, PayloadSerializer, EnumerationDataSerializer
from .models import Target, OplogEntry, Credential, EnumerationData, Payload, ExfilFile, Mitigation

# Class-based view for listing targets
@login_required
def view_oplog_exfil_file(request, pk):
    """ Serves ExfilFile - defaults to attachment """
    exfil_file = get_object_or_404(ExfilFile, pk=pk)
    try:
        file_path = exfil_file.file.path
        if not os.path.exists(file_path): raise Http404("File not found.")

        content_type, encoding = mimetypes.guess_type(file_path)
        content_type = content_type or 'application/octet-stream' # Default to download
        file = open(file_path, 'rb')
        response = HttpResponse(FileWrapper(file), content_type=content_type)
        # Default to attachment for exfil data
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
        return response
    except Exception as e:
        print(f"Error serving ExfilFile {pk}: {e}")
        raise Http404("Error accessing file.")

@login_required
def view_oplog_enum_file_inline(request, pk):
    """ Serves the SINGLE enum file from an OplogEntry record, attempting inline display. """
    oplog_entry = get_object_or_404(OplogEntry, pk=pk)
    if not oplog_entry.enum:
        raise Http404("No enum file associated with this entry.")
    try:
        file_path = oplog_entry.enum.path
        if not os.path.exists(file_path): raise Http404("File not found.")

        content_type, encoding = mimetypes.guess_type(file_path)
        content_type = content_type or 'text/plain' # Default to text
        file = open(file_path, 'rb')
        response = HttpResponse(FileWrapper(file), content_type=content_type)
        response['Content-Disposition'] = f'inline; filename="{os.path.basename(file_path)}"'
        return response
    except Exception as e:
        print(f"Error serving OplogEntry {pk} enum file: {e}")
        raise Http404("Error accessing file.")

@login_required
def view_enum_scan_file_inline(request, pk):
    """ Serves the SINGLE scan_file from an EnumerationData record, attempting inline display. """
    enum_data = get_object_or_404(EnumerationData, pk=pk)
    if not enum_data.scan_file:
        raise Http404("No scan file associated with this entry.")
    try:
        file_path = enum_data.scan_file.path
        if not os.path.exists(file_path): raise Http404("File not found.")

        content_type, encoding = mimetypes.guess_type(file_path)
        content_type = content_type or 'text/plain' # Default to text
        file = open(file_path, 'rb')
        response = HttpResponse(FileWrapper(file), content_type=content_type)
        response['Content-Disposition'] = f'inline; filename="{os.path.basename(file_path)}"'
        return response
    except Exception as e:
        print(f"Error serving EnumerationData {pk} scan_file: {e}")
        raise Http404("Error accessing file.")

class OplogEntryListCreateAPIView(generics.ListCreateAPIView):
    """
    API endpoint to list Oplog entries or create a new one.
    """
    queryset = OplogEntry.objects.all().order_by('-timestamp') # Get all entries, newest first
    serializer_class = OplogEntrySerializer
    # Require users to be authenticated to access this endpoint
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        """Automatically set the operator to the request user on create."""
        serializer.save(operator=self.request.user)

class TargetListCreateAPIView(generics.ListCreateAPIView):
    queryset = Target.objects.all()
    serializer_class = TargetSerializer
    permission_classes = [permissions.IsAuthenticated]
    # No operator to set for Target model

class CredentialListCreateAPIView(generics.ListCreateAPIView):
    queryset = Credential.objects.all()
    serializer_class = CredentialSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        # Automatically set operator on create
        serializer.save(operator=self.request.user)

class PayloadListCreateAPIView(generics.ListCreateAPIView):
    queryset = Payload.objects.all()
    serializer_class = PayloadSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        # Automatically set operator on create
        serializer.save(operator=self.request.user)

class EnumerationDataListCreateAPIView(generics.ListCreateAPIView):
    queryset = EnumerationData.objects.all()
    serializer_class = EnumerationDataSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        # Automatically set operator when creating metadata record
        serializer.save(operator=self.request.user)

    # Note: This view doesn't handle uploading the associated
    # EnumerationScanFile records via API. That requires separate endpoints
    # or more complex nested writable serializers. Users can create the
    # metadata record here and add files later via admin or future API endpoints.

@staff_member_required
def export_all_data_zip(request):
    """
    Creates a Zip archive containing CSV exports of main models and all uploaded media files.
    Outputs string representation for ForeignKeys and custom handling for OplogEntry mitigations.
    Writes to a temporary disk file first.
    """
    print("\n***** RUNNING LATEST EXPORT CODE v12 (Temp File) *****\n")

    temp_zip_file = None  # Initialize variable to hold temp file info
    temp_zip_path = None

    try:
        # --- Create a temporary file on disk ---
        # delete=False is important so we can reopen it after ZipFile closes it.
        temp_zip_file = tempfile.NamedTemporaryFile(suffix='.zip', delete=False)
        temp_zip_path = temp_zip_file.name
        temp_zip_file.close() # Close the handle, ZipFile works with the path
        print(f"Creating temporary zip archive at: {temp_zip_path}")

        # --- Write directly to the temporary file path ---
        with zipfile.ZipFile(temp_zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:

            # --- 1. Export Models as CSV ---
            models_to_export = {
                'targets': Target,
                'oplog_entries': OplogEntry,
                'credentials': Credential,
                'payloads': Payload,
                'enumeration_data': EnumerationData,
                'exfil_files': ExfilFile,
                #'mitigations': Mitigation,
            }
            print("Starting CSV Export...")

            for filename_base, model_class in models_to_export.items():
                print(f"  Exporting {model_class.__name__}")
                queryset = model_class.objects.all()

                # --- Eager load relations (same logic as before) ---
                if model_class == OplogEntry:
                     queryset = queryset.select_related('target', 'operator').prefetch_related('mitigations')
                else:
                     # ... (keep other select_related logic) ...
                     if hasattr(model_class, 'target'): queryset = queryset.select_related('target')
                     if hasattr(model_class, 'operator'): queryset = queryset.select_related('operator')
                     if hasattr(model_class, 'oplog_entry'): queryset = queryset.select_related('oplog_entry__target', 'oplog_entry__operator')


                if not queryset.exists():
                     print(f"    Skipping {model_class.__name__} - No records found.")
                     continue

                # --- Define Headers (same logic as before) ---
                if model_class == OplogEntry:
                     field_names = [ # Custom headers for OplogEntry
                         'id', 'timestamp', 'operator', 'target', 'dst_port', 'src_ip', 'src_host',
                         'src_port', 'url', 'tool', 'command', 'output', 'notes', 'sys_mod',
                         'screenshot', 'enum', 'Mitigation Names', 'Associated Findings'
                     ]
                     concrete_field_names = [f.name for f in model_class._meta.get_fields() if f.concrete and f.name != 'mitigations']
                else:
                     field_names = [f.name for f in model_class._meta.get_fields() if f.concrete]
                     concrete_field_names = field_names

                # Use StringIO for intermediate CSV creation (still efficient)
                csv_buffer = io.StringIO()
                writer = csv.writer(csv_buffer)
                writer.writerow(field_names)

                # --- Write Data Rows (same logic as before) ---
                for obj in queryset:
                     row = []
                     for field_name in concrete_field_names:
                         # ... (keep the exact same logic for getting/formatting concrete field values) ...
                         value = getattr(obj, field_name)
                         value_to_append = ''
                         try:
                             field_obj = model_class._meta.get_field(field_name)
                             if isinstance(value, datetime.datetime): value_to_append = value.isoformat()
                             elif field_obj.is_relation and not field_obj.one_to_many and not field_obj.many_to_many: value_to_append = str(value) if value is not None else ''
                             elif isinstance(value, file_fields.FieldFile): value_to_append = value.name if (value and value.name) else ''
                             elif value is None: value_to_append = ''
                             else: value_to_append = str(value)
                         except Exception as e:
                             print(f"Error processing field '{field_name}' for {model_class.__name__} PK {obj.pk}: {e}")
                             value_to_append = '[ERROR]'
                         row.append(value_to_append)

                     # --- Add custom handling for OplogEntry's M2M (same logic as before) ---
                     if model_class == OplogEntry:
                         related_mitigations = obj.mitigations.all()
                         print(f"\n--- DEBUG: Processing OplogEntry PK={obj.pk} ---")
                         print(f"Mitigation QuerySet: {related_mitigations}")
                         mitigation_names = []
                         findings_list_debug = []
                         for m in related_mitigations:
                             print(f"  - Mitigation Found: PK={m.pk}, Name='{m.name}', Finding='{m.finding}'")
                             mitigation_names.append(m.name)
                             # Add finding only if it's not blank for debug list
                             if m.finding:
                                 findings_list_debug.append(m.finding)
                         print(f"Extracted Names: {mitigation_names}")
                         print(f"Extracted Non-Blank Findings: {findings_list_debug}")
                         mitigation_names_str = ", ".join(sorted([m.name for m in related_mitigations]))
                         findings_list = [m.finding for m in related_mitigations if m.finding]
                         unique_findings_str = ", ".join(sorted(list(set(findings_list))))
                         row.append(mitigation_names_str)
                         row.append(unique_findings_str)

                     try:
                         writer.writerow(row)
                     except Exception as e: print(f"CSV Write Error for {model_class.__name__} PK {obj.pk}: {e}")

                # Write the completed CSV string to the zip file on disk
                zipf.writestr(f'{filename_base}.csv', csv_buffer.getvalue())
                # No need to close csv_buffer here, it will be garbage collected
                print(f"    Finished {filename_base}.csv")

            print("Finished CSV Export. Starting File Export...")

            # --- 2. Export Uploaded Files ---
            # (This section remains the same - writes directly to zipf on disk)
            zip_base_folder = 'uploaded_files'
            file_fields_to_export = [
                (OplogEntry, 'screenshot'), (OplogEntry, 'enum'), (ExfilFile, 'file'),
                (EnumerationData, 'scan_file'), (Payload, 'file'),
            ]
            for model_class, field_name in file_fields_to_export:
                 # ... (keep the exact same file processing logic using zipf.write) ...
                 print(f"  Processing files for {model_class.__name__}.{field_name}")
                 queryset = model_class.objects.all()
                 files_processed_count = 0
                 for obj in queryset:
                     file_field = getattr(obj, field_name)
                     if file_field and file_field.name:
                         try:
                             full_path = file_field.path
                             relative_path = file_field.name
                             zip_path = os.path.join(zip_base_folder, relative_path)
                             if os.path.exists(full_path):
                                 zipf.write(full_path, arcname=zip_path)
                                 files_processed_count += 1
                             else: print(f"Warning: File missing on disk: {full_path}")
                         except ValueError as e: print(f"Warning: ValueError accessing file path for {model_class.__name__} PK {obj.pk}, field={field_name}: {e}")
                         except Exception as e: print(f"Warning: Error adding file {getattr(file_field, 'name', 'N/A')}: {e}")
                 print(f"    Processed {files_processed_count} files for {model_class.__name__}.{field_name}")
            print("Finished File Export.")

        # --- End of 'with zipfile.ZipFile(...)' block ---
        # The temporary zip file on disk (temp_zip_path) is now complete and closed.
        print(f"Temporary zip file closed on disk: {temp_zip_path}")

        # --- 3. Prepare and Return HTTP Response using FileResponse ---
        # Reopen the completed temporary zip file for reading by FileResponse
        final_zip_file_handle = open(temp_zip_path, 'rb')

        response = FileResponse(
            final_zip_file_handle,
            as_attachment=True,
            filename=f'{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}_edc_export.zip'
        )
        print(f"Sending temporary zip file response: {temp_zip_path}")
        # FileResponse will handle streaming and closing the final_zip_file_handle

        return response

    except Exception as e:
        # Log the exception more formally if needed in production
        print(f"!!! EXPORT ERROR during zip creation or response preparation: {e}")
        # Re-raise the exception for Django's debug page during development
        raise # Or return HttpResponseServerError("An error occurred during export.")

    finally:
        # --- Clean up the temporary file from disk ---
        if temp_zip_path and os.path.exists(temp_zip_path):
            try:
                print(f"Cleaning up temporary zip file: {temp_zip_path}")
                os.remove(temp_zip_path)
            except OSError as ose:
                # Log this error, but don't prevent sending response if it was already created
                print(f"Error cleaning up temp file {temp_zip_path}: {ose}")

@staff_member_required # Ensure only staff can access
def download_sqlite_db(request):
    """ Allows staff users to download a copy of the SQLite database file. """

    db_config = settings.DATABASES.get('default', {})
    db_engine = db_config.get('ENGINE', '')

    # Ensure we are actually using SQLite
    if 'sqlite3' not in db_engine:
        return HttpResponseServerError("Database download is only configured for SQLite3.")

    db_path = db_config.get('NAME', None)

    if not db_path or not os.path.exists(db_path):
        raise Http404("Database file not found at configured path.")

    try:
        # Create a temporary file to copy the database to.
        # delete=False is important on some OSes to allow reopening by FileResponse.
        # The file will be deleted manually in the finally block.
        # Using NamedTemporaryFile ensures a unique name.
        with tempfile.NamedTemporaryFile(delete=False) as temp_db:
            temp_db_path = temp_db.name
            print(f"Copying live DB '{db_path}' to temporary file '{temp_db_path}'")
            # Copy the live database file to the temporary file
            shutil.copy2(db_path, temp_db_path) # copy2 preserves metadata if possible
            print("Copy complete.")

        # Reopen the temporary file in binary read mode for FileResponse
        # FileResponse will manage closing this file object
        final_temp_file = open(temp_db_path, 'rb')

        # Prepare the response to serve the copied file
        response = FileResponse(final_temp_file, as_attachment=True, filename=f'{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}_db_backup.sqlite3')
        print(f"Serving temporary DB file: {temp_db_path}")

        # Note: We can't reliably delete temp_db_path here because FileResponse
        # might still be streaming it. Django's FileResponse is supposed to close
        # the file object, but cleanup of the path itself might need a separate process
        # for truly temporary files in long-running views. For this admin action,
        # manual cleanup or OS temp cleaning is often sufficient. Or use a context manager
        # that cleans up AFTER the response is fully sent (more complex).

        return response

    except Exception as e:
        print(f"Error during database copy/serve: {e}")
        # Clean up temporary file if it exists and an error occurred before response
        if 'temp_db_path' in locals() and os.path.exists(temp_db_path):
             try:
                 os.remove(temp_db_path)
                 print(f"Cleaned up temporary file: {temp_db_path}")
             except OSError as ose:
                 print(f"Error cleaning up temp file {temp_db_path}: {ose}")
        return HttpResponseServerError("An error occurred during database export.")

# Classification choices
CLASSIFICATION_CHOICES = ['Unclassified', 'CUI', 'Secret', 'Secret // NOFORN']
DEFAULT_CLASSIFICATION = CLASSIFICATION_CHOICES[0] # Default to Unclassified

# Criticality prioritization choices
PRIORITY_CHOICES = ['Critical', 'High', 'Medium', 'Low', 'Informational']
# Define the sort order
PRIORITY_ORDER_MAP = {name: index for index, name in enumerate(PRIORITY_CHOICES)}
DEFAULT_PRIORITY = 'Informational' # Or choose another default

@staff_member_required # Use staff_member_required since it's linked from admin
def finding_report_view(request):
    """
    Gathers data from Oplog Entries and generates a structured report HTML page.
    """

    selected_classification = request.GET.get('classification', DEFAULT_CLASSIFICATION)
    # Validate selection against defined choices
    if selected_classification not in CLASSIFICATION_CHOICES:
        print(f"Warning: Invalid classification '{selected_classification}' received. Resetting to default.")
        selected_classification = DEFAULT_CLASSIFICATION

    print(f"DEBUG: Final selected_classification: '{selected_classification}'")

    # Data aggregation logic
    oplog_entries = OplogEntry.objects.prefetch_related(
        'mitigations', 'target'
    ).order_by('timestamp').all()

    findings_data = defaultdict(lambda: {'targets': set(), 'mitigations': set(), 'oplog_details': []})

    for entry in oplog_entries:
        target_repr = "Unknown Target"
        if entry.target:
             hostname = entry.target.hostname or "NoHostname"
             ip = entry.target.ip_address or "NoIP"
             target_repr = f"{hostname}({ip})"

        entry_mitigations = entry.mitigations.all()
        if not entry_mitigations: continue

        for mitigation in entry_mitigations:
            finding_str = mitigation.finding
            if not finding_str: continue
            data_for_finding = findings_data[finding_str]
            data_for_finding['targets'].add(target_repr)
            data_for_finding['mitigations'].add(mitigation)
            data_for_finding['oplog_details'].append({
                'id': entry.pk,
                'url': entry.url or "",
                'notes': entry.notes or "",
                'command': entry.command or "",
                'output': entry.output or "",
                'screenshot_url': entry.screenshot.url if entry.screenshot else None,
                'screenshot_path': entry.screenshot.path if entry.screenshot else None,
                'timestamp': entry.timestamp,
                'operator': entry.operator.username if entry.operator else 'Unknown',
                # Add link to the admin change view for this entry
                'admin_change_url': reverse('admin:collector_oplogentry_change', args=[entry.pk])
            })

    # Post-process aggregated data
    processed_findings = []
    all_mitigation_objects = set()
    sorted_finding_keys = sorted(findings_data.keys())

    for finding_str in sorted_finding_keys:
        data = findings_data[finding_str]
        mitigations_list = sorted(list(data['mitigations']), key=lambda m: m.name)
        all_mitigation_objects.update(mitigations_list)
        ccis = sorted(list(set(m.reference for m in mitigations_list if m.reference and m.reference.strip().upper().startswith('CCI:'))))
        targets_list = sorted(list(data['targets']))
        oplog_details_sorted = sorted(data['oplog_details'], key=lambda d: d['timestamp'])
        processed_findings.append({
            'finding_title': finding_str,
            'targets': targets_list,
            'mitigations': mitigations_list,
            'ccis': ccis,
            'oplog_details': oplog_details_sorted
        })

    summary_pairs = set()
    for finding_data in processed_findings:
        for mitigation in finding_data['mitigations']:
            summary_pairs.add((finding_data['finding_title'], mitigation.name))
    summary_table_data = sorted(list(summary_pairs), key=lambda x: (x[0], x[1]))

    print(f"Processed {len(processed_findings)} unique findings for view. Classification: {selected_classification}")

    context = {
        'report_findings': processed_findings,
        'summary_table_data': summary_table_data,
        'report_date': now().date(),
        'report_title': 'Findings Report',
        'classification_choices': CLASSIFICATION_CHOICES,
        'selected_classification': selected_classification,
        'priority_choices': PRIORITY_CHOICES,
        'default_priority': DEFAULT_PRIORITY,
        #'export_docx_url': '?export=docx' # Link for the HTML template button
        #'export_docx_url': f'?classification={selected_classification}&export=docx' #Added for classification
    }
    print(f"DEBUG: Context passed to template: selected_classification='{context['selected_classification']}', priority_choices exists={('priority_choices' in context)}, default_priority='{context['default_priority']}'") # Debug context
    print("--- finding_report_view END ---") # Debug End

    return render(request, 'collector/report_template.html', context)

@staff_member_required
@require_POST # Ensure this view only handles POST requests
def finding_report_export_docx(request):
    """
    Generates and returns the DOCX report based on POSTed priorities.
    """
    try:
        # --- Get classification and priorities from POST data ---
        selected_classification = request.POST.get('classification', DEFAULT_CLASSIFICATION)
        if selected_classification not in CLASSIFICATION_CHOICES:
            selected_classification = DEFAULT_CLASSIFICATION

        priorities_json = request.POST.get('priorities', '{}') # Default to empty JSON obj string
        try:
            # Parse the JSON mapping of finding_title -> priority
            finding_priorities_map = json.loads(priorities_json)
        except json.JSONDecodeError:
            print("Warning: Could not decode priorities JSON. Using default priority.")
            finding_priorities_map = {} # Use defaults if JSON is bad

        print(f"Generating DOCX export. Classification: {selected_classification}")
        # print(f"Received priorities map: {finding_priorities_map}") # Debug

        # --- Re-aggregate data (necessary as HTTP is stateless) ---
        # This duplicates the aggregation logic from the GET view
        oplog_entries = OplogEntry.objects.prefetch_related('mitigations', 'target').order_by('timestamp').all()
        findings_data = defaultdict(lambda: {'targets': set(), 'mitigations': set(), 'oplog_details': []})
        # ... (loop to populate findings_data exactly as in the GET view) ...
        for entry in oplog_entries:
            target_repr = "Unknown Target";
            if entry.target: hostname = entry.target.hostname or "NoHostname"; ip = entry.target.ip_address or "NoIP"; target_repr = f"{hostname}({ip})"
            entry_mitigations = entry.mitigations.all();
            if not entry_mitigations: continue
            for mitigation in entry_mitigations:
                finding_str = mitigation.finding;
                if not finding_str: continue
                data_for_finding = findings_data[finding_str]
                data_for_finding['targets'].add(target_repr); data_for_finding['mitigations'].add(mitigation)
                data_for_finding['oplog_details'].append({'id': entry.pk, 'url': entry.url or "", 'notes': entry.notes or "", 'command': entry.command or "", 'output': entry.output or "", 'screenshot_url': entry.screenshot.url if entry.screenshot else None, 'screenshot_path': entry.screenshot.path if entry.screenshot else None, 'timestamp': entry.timestamp, 'operator': entry.operator.username if entry.operator else 'Unknown', 'admin_change_url': reverse('admin:collector_oplogentry_change', args=[entry.pk])})


        # --- Prepare data structure for sorting ---
        sortable_findings = []
        for finding_str, data in findings_data.items():
            # Get priority for this finding from the map, use default if not found
            priority_name = finding_priorities_map.get(finding_str, DEFAULT_PRIORITY)
            # Get numeric sort key from priority name
            priority_sort_key = PRIORITY_ORDER_MAP.get(priority_name, len(PRIORITY_ORDER_MAP)) # Put unknowns last

            # Process details (same as in GET view)
            mitigations_list = sorted(list(data['mitigations']), key=lambda m: m.name)
            ccis = sorted(list(set(m.reference for m in mitigations_list if m.reference and m.reference.strip().upper().startswith('CCI:'))))
            targets_list = sorted(list(data['targets']))
            oplog_details_sorted = sorted(data['oplog_details'], key=lambda d: d['timestamp'])

            sortable_findings.append({
                'priority_sort_key': priority_sort_key,
                'priority_name': priority_name, # Keep name if needed later
                'finding_title': finding_str,
                'targets': targets_list,
                'mitigations': mitigations_list,
                'ccis': ccis,
                'oplog_details': oplog_details_sorted
            })

        # --- Sort findings by priority, then alphabetically ---
        sorted_report_findings = sorted(sortable_findings, key=lambda x: (x['priority_sort_key'], x['finding_title']))
        print(f"Sorted {len(sorted_report_findings)} findings for DOCX export.")

        # --- Prepare sorted summary data ---
        summary_table_data_sorted = []
        for finding_data in sorted_report_findings:
             for mitigation in finding_data['mitigations']:
                  summary_table_data_sorted.append((finding_data['finding_title'], mitigation.name))
        # summary_table_data_sorted is now ordered by priority then finding then mitigation

        # --- Generate DOCX Document ---
        document = docx.Document()
        # Add Footer with classification
        section = document.sections[0]; footer = section.footer; header = section.header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.text = selected_classification; header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.text = selected_classification; footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add main content
        document.add_heading('Findings Report', level=1)
        document.add_paragraph(f"Report Generated: {now().date()}")
        document.add_paragraph()

        # Loop through the *SORTED* findings
        for i, finding_data in enumerate(sorted_report_findings, 1): # Start numbering at 1
            # Use the new counter 'i' for the heading number
            document.add_heading(f"Finding {i}: {finding_data['finding_title']}", level=2)

            # Create and populate the main table (using same logic as before, just with finding_data from sorted list)
            table = document.add_table(rows=0, cols=2, style='Table Grid'); table.autofit = False; table.allow_autofit = False
            table.columns[0].width = Inches(2.0); table.columns[1].width = Inches(5.0)
            def add_row(label, value): row_cells = table.add_row().cells; row_cells[0].text = label; row_cells[1].text = str(value) if value is not None else ''

            # Populate rows, using the current priority for the first row!
            add_row("Mitigation Priority", f"{finding_data['finding_title']}\n[Priority Set: {finding_data['priority_name']}]") # Show selected priority
            add_row("Description", "[User to provide detailed description...]")
            add_row("Affected Resources", ", ".join(finding_data['targets']) if finding_data['targets'] else 'N/A')
            add_row("Operational Impact", "[User to describe operational impact...]")
            add_row("Threat Posture", "[User to describe threat posture...]")
            # ... (Mitigation(s) Row - keep corrected logic from previous step) ...
            mitigation_cells = table.add_row().cells; mitigation_cells[0].text = "Mitigation(s)"; mitigation_cells[1].text = ""
            if finding_data['mitigations']:
                for mitigation in finding_data['mitigations']: p_name = mitigation_cells[1].add_paragraph(); p_name.add_run(mitigation.name).bold = True; p_desc = mitigation_cells[1].add_paragraph(mitigation.description); p_desc.paragraph_format.space_after = Pt(6)
            else: mitigation_cells[1].add_paragraph("N/A")
            add_row("Control Correlation Identifier (CCI)", ", ".join(finding_data['ccis']) if finding_data['ccis'] else 'N/A')
            add_row("CVSS Score", "[User to provide CVSS Score...]")
            # ... (PoC Heading Row - keep logic) ...
            poc_heading_cells = table.add_row().cells; merged_poc_heading = poc_heading_cells[0].merge(poc_heading_cells[1]); para = merged_poc_heading.paragraphs[0]; para.text = ""; run = para.add_run("Proof of Concept"); run.bold = True; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # ... (PoC Details Row - keep logic, iterating finding_data['oplog_details']) ...
            poc_detail_cells = table.add_row().cells; merged_poc_details = poc_detail_cells[0].merge(poc_detail_cells[1]); merged_poc_details.text = ""
            if finding_data['oplog_details']:
                 for detail in finding_data['oplog_details']:
                     p = merged_poc_details.add_paragraph(); p.add_run(f"Entry {detail['id']} ({detail['timestamp'].strftime('%Y-%m-%d %H:%i')} by {detail['operator']}):").bold = True
                     if detail.get('url'): merged_poc_details.add_paragraph(f"URL: {detail['url']}") # etc... copy rest of PoC detail logic
                     if detail.get('notes'): merged_poc_details.add_paragraph(f"Notes:\n{detail['notes']}")
                     if detail.get('command'): p = merged_poc_details.add_paragraph("Command:"); p_code = merged_poc_details.add_paragraph(detail['command']); p_code.style = 'Normal'; p_code.runs[0].font.name = 'Courier New'; p_code.paragraph_format.left_indent = Inches(0.25)
                     if detail.get('output'): p = merged_poc_details.add_paragraph("Output:"); p_code = merged_poc_details.add_paragraph(detail['output']); p_code.style = 'Normal'; p_code.runs[0].font.name = 'Courier New'; p_code.paragraph_format.left_indent = Inches(0.25)
                     if detail.get('screenshot_path') and os.path.exists(detail['screenshot_path']):
                         try: merged_poc_details.add_paragraph("Screenshot:"); document.add_picture(detail['screenshot_path'], width=Inches(5.0))
                         except Exception as img_e: print(f"Error adding screenshot {detail['screenshot_path']}: {img_e}"); merged_poc_details.add_paragraph(f"[Error adding screenshot: {os.path.basename(detail['screenshot_path'])}]")
                     elif detail.get('screenshot_url'): merged_poc_details.add_paragraph(f"Screenshot URL: {detail['screenshot_url']} (embedding requires URL fetch)")
                     if detail != finding_data['oplog_details'][-1]: merged_poc_details.add_paragraph("---")
            else: merged_poc_details.add_paragraph("No specific Oplog entry details linked to this finding.")


        # --- Add *SORTED* Summary Table ---
        if summary_table_data_sorted: # Use the sorted data
            document.add_heading("Mitigation Priorities Summary", level=2)
            summary_table = document.add_table(rows=1, cols=3, style='Table Grid')
            summary_table.autofit = True
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Finding'; hdr_cells[1].text = 'Mitigation Priority'; hdr_cells[2].text = 'Mitigation'
            # Add sorted data rows
            for finding_title, mitigation_name in summary_table_data_sorted:
                row_cells = summary_table.add_row().cells
                row_cells[0].text = finding_title
                # Get priority for this finding again to display in summary (or retrieve from sorted_report_findings)
                priority_display = finding_priorities_map.get(finding_title, DEFAULT_PRIORITY)
                row_cells[1].text = priority_display # Display selected priority
                row_cells[2].text = mitigation_name

        # Save and return response (same as before)
        buffer = io.BytesIO(); document.save(buffer); buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        timestamp = now().strftime('%Y%m%d_%H%M%S'); response['Content-Disposition'] = f'attachment; filename="{timestamp}_findings_report_{selected_classification.replace(" ","_").replace("/","_")}.docx"'
        print("DOCX export prepared successfully.")
        return response

    except Exception as e:
        print(f"!!! ERROR generating DOCX report: {e}")
        # import traceback; traceback.print_exc();
        return HttpResponse(f"Error generating Word report: {e}", status=500)
